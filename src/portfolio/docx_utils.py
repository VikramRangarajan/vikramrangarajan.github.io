from docx.oxml.parser import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.text.run import Run
from docx.text.paragraph import Paragraph
from docx.table import _Cell
from docx.shared import Cm, RGBColor, Pt
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.document import Document


ITALIC = 0
BOLD = 1
UNDERLINE = 2


def add_run(p: Paragraph, text: str, font_size: Pt, *args) -> Run:
    r = p.add_run(text)
    r.font.size = font_size
    for arg in args:
        if arg == ITALIC:
            r.italic = True
        if arg == BOLD:
            r.bold = True
        if arg == UNDERLINE:
            r.underline = True
    return r


def set_global_font(document: Document, font_name):
    # Set font for paragraphs
    for paragraph in document.paragraphs:
        # Set font for hyperlinks
        for hl in paragraph.hyperlinks:
            for run in hl.runs:
                run.font.name = font_name


        for run in paragraph.runs:
            run.font.name = font_name

    # Set font for tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name



def insertHR(paragraph: Paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")

    pPr.insert_element_before(
        pBdr,
        "w:shd",
        "w:tabs",
        "w:suppressAutoHyphens",
        "w:kinsoku",
        "w:wordWrap",
        "w:overflowPunct",
        "w:topLinePunct",
        "w:autoSpaceDE",
        "w:autoSpaceDN",
        "w:bidi",
        "w:adjustRightInd",
        "w:snapToGrid",
        "w:spacing",
        "w:ind",
        "w:contextualSpacing",
        "w:mirrorIndents",
        "w:suppressOverlap",
        "w:jc",
        "w:textDirection",
        "w:textAlignment",
        "w:textboxTightWrap",
        "w:outlineLvl",
        "w:divId",
        "w:cnfStyle",
        "w:rPr",
        "w:sectPr",
        "w:pPrChange",
    )
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)


def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ("start", "top", "end", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = "w:{}".format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag), None)
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn("w:{}".format(key)), str(edge_data[key]))


def set_table_border_none(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"sz": 0, "val": "single", "color": "#FFFFFF"},
                bottom={"sz": 0, "val": "single", "color": "#FFFFFF"},
                start={"sz": 0, "val": "single", "color": "#FFFFFF"},
                end={"sz": 0, "val": "single", "color": "#FFFFFF"},
            )


def get_usable_width(doc):
    section = doc.sections[0]
    page_width = section.page_width
    left_margin = section.left_margin
    right_margin = section.right_margin

    page_width_inches = page_width.inches
    left_margin_inches = left_margin.inches
    right_margin_inches = right_margin.inches

    # Calculate usable width
    usable_width_inches = page_width_inches - (left_margin_inches + right_margin_inches)
    return usable_width_inches


def set_column_widths(table, widths):
    for col, width in zip(table.columns, widths):
        col.width = Cm(width)


def right_align_columns(table, col_idxs=(1,)):
    for idx in col_idxs:
        for row in table.rows:
            for paragraph in row.cells[idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_hyperlink(paragraph, text, url, font_size):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = Run(OxmlElement("w:r"), paragraph)
    new_run.font.size = font_size
    new_run.text = text

    # Set the run's style to the builtin hyperlink style, defining it if necessary
    new_run.style = get_or_create_hyperlink_style(part.document)

    # Join all the xml elements together
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)
    paragraph._p.r_lst.append(new_run)
    return hyperlink, new_run


def get_or_create_hyperlink_style(d):
    """If this document had no hyperlinks so far, the builtin
    Hyperlink style will likely be missing and we need to add it.
    There's no predefined value, different Word versions
    define it differently.
    This version is how Word 2019 defines it in the
    default theme, excluding a theme reference.
    """
    if "Hyperlink" not in d.styles:
        if "Default Character Font" not in d.styles:
            ds = d.styles.add_style(
                "Default Character Font", WD_STYLE_TYPE.CHARACTER, True
            )
            ds.element.set(qn("w:default"), "1")
            ds.priority = 1
            ds.hidden = True
            ds.unhide_when_used = True
            del ds
        hs = d.styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER, True)
        hs.base_style = d.styles["Default Character Font"]
        hs.unhide_when_used = True
        hs.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
        hs.font.underline = True
        del hs

    return "Hyperlink"
