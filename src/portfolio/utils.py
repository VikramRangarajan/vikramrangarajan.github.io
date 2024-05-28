from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.run import Run
from docx.shared import Cm
from docx.opc.constants import RELATIONSHIP_TYPE
import docx


class ddict(dict):
    __getattr__ = dict.__getitem__
    __setattr__ = dict.__setitem__
    __delattr__ = dict.__delitem__
    __str__ = dict.__str__

    def __init__(self, dct):
        for key, value in dct.items():
            if hasattr(value, "keys"):
                value = ddict(value)
            self[key] = value


def set_global_font(document, font_name):
    # Set font for paragraphs
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name

    # Set font for tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name


def insertHR(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    pPr.insert_element_before(
        pBdr,
        "w:shd",
        "w:tabs",
        "w:suppressAutoHyphens",
        "w:kinsoku",
        "w:wordWrap",
        "w:overflowPunct",
        "w:topLinePunct",
        "w:autoSpaceDE",
        "w:autoSpaceDN",
        "w:bidi",
        "w:adjustRightInd",
        "w:snapToGrid",
        "w:spacing",
        "w:ind",
        "w:contextualSpacing",
        "w:mirrorIndents",
        "w:suppressOverlap",
        "w:jc",
        "w:textDirection",
        "w:textAlignment",
        "w:textboxTightWrap",
        "w:outlineLvl",
        "w:divId",
        "w:cnfStyle",
        "w:rPr",
        "w:sectPr",
        "w:pPrChange",
    )
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)


def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ("start", "top", "end", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = "w:{}".format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn("w:{}".format(key)), str(edge_data[key]))


def set_table_border_none(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"sz": 0, "val": "single", "color": "#FFFFFF"},
                bottom={"sz": 0, "val": "single", "color": "#FFFFFF"},
                start={"sz": 0, "val": "single", "color": "#FFFFFF"},
                end={"sz": 0, "val": "single", "color": "#FFFFFF"},
            )


def get_usable_width(doc):
    section = doc.sections[0]
    page_width = section.page_width
    left_margin = section.left_margin
    right_margin = section.right_margin

    page_width_cm = page_width / 360000
    left_margin_cm = left_margin / 360000
    right_margin_cm = right_margin / 360000

    # Calculate usable width
    usable_width_cm = page_width_cm - (left_margin_cm + right_margin_cm)
    return usable_width_cm


def set_column_widths(table, widths):
    for col, width in zip(table.columns, widths):
        col.width = Cm(width)


def add_date_table(doc, rows):
    table = doc.add_table(rows=rows, cols=2)
    set_table_border_none(table)
    width = get_usable_width(doc)
    set_column_widths(table, (width * 2 / 3, width * 1 / 3))
    return table


def right_align_columns(table, col_idxs=(1,)):
    for idx in col_idxs:
        for row in table.rows:
            for paragraph in row.cells[idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(
        qn("r:id"),
        r_id,
    )

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = Run(OxmlElement("w:r"), paragraph)
    new_run.text = text

    # Set the run's style to the builtin hyperlink style, defining it if necessary
    new_run.style = get_or_create_hyperlink_style(part.document)

    # Join all the xml elements together
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)
    return hyperlink


def get_or_create_hyperlink_style(d):
    """If this document had no hyperlinks so far, the builtin
    Hyperlink style will likely be missing and we need to add it.
    There's no predefined value, different Word versions
    define it differently.
    This version is how Word 2019 defines it in the
    default theme, excluding a theme reference.
    """
    if "Hyperlink" not in d.styles:
        if "Default Character Font" not in d.styles:
            ds = d.styles.add_style(
                "Default Character Font", docx.enum.style.WD_STYLE_TYPE.CHARACTER, True
            )
            ds.element.set(docx.oxml.shared.qn("w:default"), "1")
            ds.priority = 1
            ds.hidden = True
            ds.unhide_when_used = True
            del ds
        hs = d.styles.add_style(
            "Hyperlink", docx.enum.style.WD_STYLE_TYPE.CHARACTER, True
        )
        hs.base_style = d.styles["Default Character Font"]
        hs.unhide_when_used = True
        hs.font.color.rgb = docx.shared.RGBColor(0x05, 0x63, 0xC1)
        hs.font.underline = True
        del hs

    return "Hyperlink"
