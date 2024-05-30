import json
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx_utils import (
    insertHR,
    set_global_font,
    ddict,
    get_usable_width,
    add_hyperlink,
)

def create_docx():
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    ROOT = Path(__file__).parent.absolute()


    with open(ROOT / "portfolio.json") as f:
        por = ddict(json.load(f))

    width = get_usable_width(doc)

    # Add Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(por.current.info.name)
    r.font.size = Pt(24)

    # Info and Links section at the top
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Cm(width), WD_TAB_ALIGNMENT.RIGHT)
    email = por.current.info.email
    phone = por.current.info.phone
    linkedin = por.current.info.linkedin
    github = por.current.info.github
    website = por.current.info.website
    p.add_run("Website: ")
    add_hyperlink(p, website, website)
    p.add_run("\t")
    p.add_run("Email: ")
    add_hyperlink(p, f"{email}\n", f"mailto:{email}")
    p.add_run("LinkedIn: ")
    add_hyperlink(p, linkedin, linkedin)
    p.add_run(f"\tLocation: {por.current.info.location}\n")
    p.add_run("GitHub: ")
    add_hyperlink(p, github, github)
    p.add_run("\t")
    p.add_run("Phone: ")
    add_hyperlink(p, phone, f"tel:{phone}")

    # Add education section
    p = doc.add_paragraph()
    p.add_run("Education").font.size = Pt(16)
    insertHR(p)
    p = doc.add_paragraph()
    college = por.current.education[0]
    p.add_run(
        college["degree"]
        + " -- "
        + college["major"]
        + (", " + college["minor"] + " Minor" if "minor" in college.keys() else "")
        + "\n"
    ).bold = True
    p.add_run(college["name"] + ", " + college["location"] + "\n")

    if college["expected"] is not None:
        p.add_run(
            college["start"] + " - Expected " + college["expected"] + "\n"
        ).italic = True
    else:
        p.add_run(college["start"] + " - " + college["end"] + "\n").italic = True
    p.add_run("GPA: " + str(college["GPA"]) + "\n")
    p.add_run("Relevant Coursework: ").italic = True
    p.add_run(", ".join(college["coursework"]))

    # Experience Section
    p = doc.add_paragraph()
    p.add_run("Experience & Projects").font.size = Pt(16)
    insertHR(p)

    experiences = por.current.experience
    for exp in experiences:
        p = doc.add_paragraph()
        l1 = exp["company-name"]
        if exp["location"] is not None:
            l1 += ", " + exp["location"]
        l1 += "\n"
        if exp["title_link"] is not None:
            _, j = add_hyperlink(p, l1, exp["title_link"])
        else:
            j = p.add_run(l1)
        j.bold = j.underline = True
        if exp["title"] is not None:
            p.add_run(exp["title"] + "\n")
        time = exp["start"] + " - "
        if exp["end"] is None:
            time += "Present"
        else:
            time += exp["end"]
        p.add_run(time).italic = True
        for line in exp["description"].split("\n"):
            doc.add_paragraph(line, "List Bullet")

    # Skills Section
    p = doc.add_paragraph()
    p.add_run("Technical Skills").font.size = Pt(16)
    insertHR(p)

    p = doc.add_paragraph()
    p.add_run("Programming Languages: ").bold = True
    p.add_run(", ".join(por.current.skills["programming-languages"]))
    p.add_run("\nTechnologies: ").bold = True
    p.add_run(", ".join(por.current.skills.technologies))


    # Certs Section
    p = doc.add_paragraph()
    p.add_run("Awards & Certifications").font.size = Pt(16)
    insertHR(p)

    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Cm(width), WD_TAB_ALIGNMENT.RIGHT)
    for i, award in enumerate(por.current.awards):
        p.add_run(f"{award['name']}\t").bold = True
        p.add_run(f"{award['date']}\n").italic = True

    set_global_font(doc, "Calibri")
    doc.save("formatted_document.docx")

if __name__ == "__main__":
    create_docx()