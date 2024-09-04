from pathlib import Path
from subprocess import run
from datetime import date
import re

from docx.document import Document
from docx import Document as CreateDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT

from data import Portfolio, Education, str_date, DocxConfig, Margins
from docx_utils import (
    insertHR,
    set_global_font,
    get_usable_width,
    add_hyperlink,
    add_run,
    ITALIC,
    BOLD,
)


def add_education(doc: Document, education: Education, docx_config: DocxConfig):
    if education.current is False:
        return
    p = doc.add_paragraph()
    text_size = docx_config.text_font_size
    if education.degree is None:
        raise ValueError("No Degree Specified for Education")
    if education.major is None:
        raise ValueError("No Major Specified for Education")
    if education.coursework is None:
        raise ValueError("No Coursework Specified for Education")
    minor = ", " + education.minor + " Minor" if education.minor is not None else ""

    add_run(
        p,
        f"{education.degree} -- {education.major}{minor}\n",
        Pt(text_size),
        BOLD,
    )
    add_run(p, education.name + ", " + education.location + "\n", Pt(text_size))

    if education.expected is not None:
        add_run(
            p,
            f"{str_date(education.start)} - Expected {str_date(education.expected)}\n",
            Pt(text_size),
            ITALIC,
        )
    elif education.end is not None:
        add_run(
            p,
            f"{str_date(education.start)} - {str_date(education.end)}\n",
            Pt(text_size),
            ITALIC,
        )
    add_run(p, f"GPA: {str(education.GPA)}\n", Pt(text_size))
    add_run(p, "Relevant Coursework: ", Pt(text_size), ITALIC)
    add_run(p, ", ".join(education.coursework), Pt(text_size))


def add_info_and_links(doc: Document, por: Portfolio, docx_config: DocxConfig):
    text_size = docx_config.text_font_size
    p = doc.add_paragraph()
    width = get_usable_width(doc)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(width), WD_TAB_ALIGNMENT.RIGHT)
    email = por.info.email
    phone = por.info.phone
    linkedin = por.info.linkedin
    github = por.info.github
    website = por.info.website
    add_run(p, "Website: ", Pt(text_size))
    add_hyperlink(p, website, website, Pt(text_size))
    add_run(p, "\tEmail: ", Pt(text_size))
    add_hyperlink(p, f"{email}\n", f"mailto:{email}", Pt(text_size))
    add_run(p, "LinkedIn: ", Pt(text_size))
    add_hyperlink(p, linkedin, linkedin, Pt(text_size))
    add_run(p, f"\tLocation: {por.info.location}\n", Pt(text_size))

    add_run(p, "GitHub: ", Pt(text_size))
    add_hyperlink(p, github, github, Pt(text_size))
    add_run(p, "\tPhone: ", Pt(text_size))
    add_hyperlink(p, phone, f"tel:{phone}", Pt(text_size))


def add_experiences(doc: Document, por: Portfolio, docx_config: DocxConfig):
    text_size = docx_config.text_font_size
    p = doc.add_paragraph()
    p.add_run("Experience & Projects").font.size = Pt(docx_config.heading_font_size)
    insertHR(p)

    experiences = por.experience
    if experiences is not None:
        for exp in experiences:
            if exp.current is False:
                continue
            p = doc.add_paragraph()
            l1 = exp.company_name
            if exp.location is not None:
                l1 += ", " + exp.location
            l1 += "\n"
            if exp.title_link is not None:
                _, j = add_hyperlink(p, l1, exp.title_link, Pt(text_size))
            else:
                j = p.add_run(l1)
            j.font.size = Pt(text_size)
            j.bold = j.underline = True
            if exp.title is not None:
                add_run(p, exp.title + "\n", Pt(text_size))
            time = str_date(exp.start) + " - "
            if exp.end is None:
                time += "Present"
            else:
                time += str_date(exp.end)
            add_run(p, time, Pt(text_size), ITALIC)
            for line in exp.description:
                doc.add_paragraph(line, "List Bullet")


def add_publications(doc: Document, por: Portfolio, docx_config: DocxConfig):
    text_size = docx_config.text_font_size
    p = doc.add_paragraph()
    add_run(p, "Publications", Pt(docx_config.heading_font_size))
    insertHR(p)

    if por.publications is not None:
        for pub in por.publications:
            p = doc.add_paragraph("", "List Number")
            for i, author in enumerate(pub.authors):
                if i < len(pub.authors) - 1:
                    delimiter = ", "
                else:
                    delimiter = "\n"
                r = add_run(p, f"{author}{delimiter}", Pt(text_size))
                if author == por.info.name:
                    r.bold = True
            pub_str = pub.title
            if "published" not in pub.status.lower():
                pub_str += f"\n{pub.status}"
            add_run(p, pub_str, Pt(text_size))


def add_skills(doc: Document, por: Portfolio, docx_config: DocxConfig):
    text_size = docx_config.text_font_size
    p = doc.add_paragraph()
    add_run(p, "Technical Skills", Pt(docx_config.heading_font_size))
    insertHR(p)

    p = doc.add_paragraph()
    if por.skills is not None:
        first = True
        for name in por.skills.keys():
            if not first:
                p.add_run("\n")
            skill_group = por.skills[name]
            name: str = re.sub("[^0-9a-zA-Z]+", " ", name).capitalize()
            name = f"{name}: "
            add_run(p, name, Pt(text_size), BOLD)
            add_run(p, ", ".join(skill_group), Pt(text_size))
            first = False


def add_awards(doc: Document, por: Portfolio, docx_config: DocxConfig):
    text_size = docx_config.text_font_size
    width = get_usable_width(doc)
    p = doc.add_paragraph()
    add_run(p, "Awards & Certifications", Pt(docx_config.heading_font_size))
    insertHR(p)

    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Inches(width), WD_TAB_ALIGNMENT.RIGHT)
    if por.awards is not None:
        for award in por.awards:
            if award.current is False:
                continue
            add_run(p, f"{award.name}\t", Pt(text_size), BOLD)
            if isinstance(award.date, date):
                date_str = str_date(award.date)
            else:
                date_str = award.date
            add_run(p, f"{date_str}\n", Pt(text_size), ITALIC)


def create_docx(path: str | Path, docx_config: DocxConfig):
    doc = CreateDocument()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(docx_config.margins.top)
        section.bottom_margin = Inches(docx_config.margins.bottom)
        section.left_margin = Inches(docx_config.margins.left)
        section.right_margin = Inches(docx_config.margins.right)

    ROOT = Path(__file__).parent.absolute()

    with open(ROOT / "portfolio.json") as f:
        por = Portfolio.model_validate_json(f.read())

    # Add Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(por.info.name)
    r.font.size = Pt(docx_config.title_font_size)

    # Info and Links section at the top
    add_info_and_links(doc, por, docx_config)

    # Add education section
    p = doc.add_paragraph()
    p.add_run("Education").font.size = Pt(docx_config.heading_font_size)
    insertHR(p)
    if por.education is not None:
        add_education(doc, por.education[0], docx_config)

    # Experience Section
    add_experiences(doc, por, docx_config)

    # Publications Section
    add_publications(doc, por, docx_config)

    # Skills Section
    add_skills(doc, por, docx_config)

    # Certs Section
    add_awards(doc, por, docx_config)

    set_global_font(doc, "Calibri")
    doc.save(str(path))


def docx_to_pdf(docx_file: str | Path):
    run(
        [
            "libreoffice",
            "--headless",
            "--convert-to",
            "pdf",
            str(docx_file),
            "--outdir",
            str(Path(docx_file).parent),
        ]
    )


if __name__ == "__main__":
    ROOT = Path(__file__).parent.parent.parent.absolute()
    docx_config = DocxConfig(
        margins=Margins(top=0.75, bottom=0.75, left=0.75, right=0.75),
        title_font_size=20,
        heading_font_size=14,
        text_font_size=11,
    )
    create_docx(ROOT / "docs/source/_static/resume.docx", docx_config)
    docx_to_pdf(ROOT / "docs/source/_static/resume.docx")
